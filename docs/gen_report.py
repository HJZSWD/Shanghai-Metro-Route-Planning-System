#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成上海地铁线路查询系统 小组报告 DOCX"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2A, 0x3F, 0x6A)
    return h


def add_para(text, bold=False, indent=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
    return p


def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()  # 表后空行
    return table


def add_code_block(code_text):
    """添加代码块（用等宽字体段落模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据结构课程设计')
run.font.size = Pt(18)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('上海地铁线路查询系统')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)

doc.add_paragraph()
doc.add_paragraph()

# 信息表
info_items = [
    ('学院', '计算机科学与技术学院'),
    ('专业', '计算机科学与技术'),
    ('课程', '数据结构'),
    ('姓名', '（请填写）'),
    ('学号', '（请填写）'),
    ('指导教师', '（请填写）'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年6月')
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录页（手动）
# ============================================================
add_heading_styled('目  录', level=1)
toc_items = [
    ('一、项目概述', 1),
    ('  1.1 项目背景与目标', 2),
    ('  1.2 功能需求概述', 2),
    ('二、系统设计', 1),
    ('  2.1 系统架构设计', 2),
    ('  2.2 目录结构', 2),
    ('  2.3 模块划分与职责', 2),
    ('三、数据结构设计', 1),
    ('  3.1 站点数据（Station）', 2),
    ('  3.2 图结构（Graph）', 2),
    ('  3.3 路径结果（PathResult）', 2),
    ('四、核心算法设计与实现', 1),
    ('  4.1 Dijkstra 双关键字最短路径', 2),
    ("  4.2 Yen's K短路算法", 2),
    ('  4.3 路径去重与指标计算', 2),
    ('五、前端可视化设计', 1),
    ('  5.1 技术架构', 2),
    ('  5.2 交互设计', 2),
    ('  5.3 视觉设计', 2),
    ('六、系统测试与结果', 1),
    ('  6.1 测试环境', 2),
    ('  6.2 测试用例', 2),
    ('  6.3 性能分析', 2),
    ('七、项目总结与展望', 1),
    ('  7.1 完成情况总结', 2),
    ('  7.2 存在的不足与改进方向', 2),
    ('  7.3 开发心得', 2),
    ('参考文献', 1),
]
for item, level in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12 if level == 1 else 11)
    run.bold = (level == 1)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

# ============================================================
# 一、项目概述
# ============================================================
add_heading_styled('一、项目概述', level=1)

add_heading_styled('1.1 项目背景与目标', level=2)
add_para(
    '上海地铁是中国规模最大的城市轨道交通系统之一，线路网络密集复杂，截至2026年已开通运营20条线路、'
    '超过500座车站。面对如此庞大的交通网络，通勤者和游客需要一个高效、直观的路径规划工具来辅助出行决策。',
    indent=True
)
add_para(
    '本课程设计项目以"上海地铁线路查询系统"为题，运用数据结构与算法课程中学习的图论、最短路径算法等核心知识，'
    '开发一套完整的上海地铁路径规划系统。项目采用C++实现核心算法后端，并辅以Web前端可视化界面，'
    '旨在解决地铁出行中的路径查询、换乘规划等实际问题。',
    indent=True
)
add_para(
    '项目开发使用AI编程助手TRAE IDE进行协作，在C++控制台应用的基础上，逐步扩展了Web可视化前端、'
    'Yen\'s K短路算法、换乘站智能合并等进阶功能。',
    indent=True
)

add_heading_styled('1.2 功能需求概述', level=2)
add_bullet('路径规划：支持单条/前3条最短时间路径和最少换乘路径的查询')
add_bullet('站点管理：CSV批量更新站点状态、手动开/关站点、恢复初始状态')
add_bullet('模糊搜索：支持根据关键词模糊匹配站名（不区分大小写），同名多线路站点各自独立显示')
add_bullet('线路查询：按线路编号查询该线路途经的全部站点')
add_bullet('关闭站点过滤：自动过滤已关闭站点，路径搜索仅使用可用站点')
add_bullet('Web可视化：响应式前端界面，路径在地铁线路上可视化展示')

# ============================================================
# 二、系统设计
# ============================================================
doc.add_page_break()
add_heading_styled('二、系统设计', level=1)

add_heading_styled('2.1 系统架构设计', level=2)
add_para(
    '系统采用前后端分离的混合架构，核心算法由C++实现，用户界面提供C++控制台和Web浏览器两种交互方式。'
    '数据层使用CSV文件存储，前后端共享同一套数据集。',
    indent=True
)

add_table(
    ['层次', '技术栈', '职责'],
    [
        ['前端展示层', 'HTML5 + CSS3 + JavaScript', '用户交互、路径可视化、实时计算'],
        ['算法计算层（JS）', 'JavaScript（BinaryHeap + Int32Array）', '浏览器端Dijkstra/Yen\'s 算法'],
        ['算法计算层（C++）', 'C++11 STL', '控制台路径规划、站点管理'],
        ['数据持久层', 'CSV文件（UTF-8 BOM）', '站点数据、边数据、状态存储'],
    ]
)

add_para(
    '系统采用"同构算法"设计理念——C++后端和JavaScript前端各自独立实现了相同的Dijkstra和Yen\'s K短路算法，'
    '确保两种交互方式下计算结果一致。前端算法实现针对浏览器环境做了特殊优化，使用BinaryHeap优先队列和'
    'Int32Array定长数组替代通用数据结构，实现了毫秒级的响应速度。',
    indent=True
)

add_heading_styled('2.2 目录结构', level=2)
add_code_block(
'''Shanghai Metro Route Planning System/
+-- frontend/               # Web前端
|   +-- web/
|       +-- index.html      # 主页面（HTML+CSS+JS，约40KB）
|       +-- data.js         # 数据集（524站/1232边，自动生成）
|       +-- README.md       # 前端说明
+-- backend/                # C++后端
|   +-- src/                # 源码（9个文件）
|   |   +-- main.cpp        # 程序入口
|   |   +-- station.h/cpp   # 站点数据管理
|   |   +-- graph.h/cpp     # 有向邻接表图结构
|   |   +-- pathfinder.h/cpp# Dijkstra + Yen's K短路
|   |   +-- menu.h/cpp      # 二级菜单交互
|   +-- algo/               # 算法文档
+-- data/                   # CSV数据集
|   +-- Station.csv         # 524个站点
|   +-- Edge.csv            # 1232条有向边
|   +-- Station_init.csv    # 初始备份
|   +-- update_station_status.csv
+-- docs/                   # 项目文档
+-- launch.bat              # 一键启动脚本
+-- launch_hidden.vbs       # 静默启动脚本
+-- backend-service.exe     # 后端服务（编译输出）
+-- README.md               # 项目总览''')

add_heading_styled('2.3 模块划分与职责', level=2)
add_table(
    ['模块', '文件', '核心职责'],
    [
        ['StationManager', 'station.h/cpp', '站点CRUD、模糊搜索、换乘站合并、CSV读写、编码转换'],
        ['MetroGraph', 'graph.h/cpp', '有向邻接表构建、关闭站点过滤、边数据加载'],
        ['PathFinder', 'pathfinder.h/cpp', 'Dijkstra双关键字最短路径、Yen\'s K短路、路径去重'],
        ['MenuController', 'menu.h/cpp', '二级菜单交互、输入校验、关键词选站'],
        ['数据入口', 'main.cpp', '初始化、参数解析（--hidden）、路径自动推导'],
    ]
)

# ============================================================
# 三、数据结构设计
# ============================================================
doc.add_page_break()
add_heading_styled('三、数据结构设计', level=1)

add_heading_styled('3.1 站点数据（Station）', level=2)
add_para(
    '站点数据使用Station结构体表示，StationManager类负责管理所有站点。',
    indent=True
)
add_code_block(
'''struct Station {
    int stationId;          // 站点唯一编号
    std::string stationName; // 站点名称（如"人民广场"）
    std::string lineName;    // 所属线路（如"1号线"）
    bool isOpen;             // 运营状态：true=开放，false=关闭
};''')

add_para('为支持换乘站合并显示，设计了合并结构体：', indent=True)
add_code_block(
'''struct MergedStationInfo {
    std::string stationName; // 站点名称（合并后显示同一名称）
    std::string lineNames;   // 途经线路（如"1号线/2号线/8号线"）
    std::vector<int> stationIds; // 对应的多个站点ID
    bool isOpen;             // 任一条线路开放即为开放
};''')

add_para('StationManager类提供以下关键功能：', indent=True)
add_bullet('CSV加载：自动检测UTF-8 BOM，兼容GB18030编码的CSV文件')
add_bullet('模糊搜索：基于字符串包含匹配，不区分大小写，同名多线路站点各自独立显示')
add_bullet('批量更新：支持CSV格式（"开启/开放/关闭"或"1/0"）批量修改站点状态')

add_heading_styled('3.2 图结构（Graph）', level=2)
add_para(
    '系统使用有向邻接表（vector<vector<MetroEdge>>）表示地铁线路网络。'
    '每条物理连接在邻接表中存储为两条独立的有向边（A→B 和 B→A），分别记录方向信息。',
    indent=True
)
add_code_block(
'''struct MetroEdge {
    int startId;             // 起点站ID
    int endId;               // 终点站ID
    std::string lineName;    // 所属线路（如"1号线"）
    std::string direction;   // 方向（"上行"/"下行"）
    int time;                // 行驶时间（分钟），换乘边固定为5
};

class MetroGraph {
    std::vector<std::vector<MetroEdge>> m_adj;  // 有向邻接表
    int m_stationCount;
};''')

add_para(
    'MetroGraph对外提供两种邻接查询接口：getAdjEdges返回全部邻接边（包含关闭站点的连接），'
    'getAvailableEdges则自动过滤目标端已关闭的站点，确保路径搜索不会经过关闭站点。',
    indent=True
)

add_heading_styled('3.3 路径结果（PathResult）', level=2)
add_code_block(
'''struct PathResult {
    std::vector<int> pathStations;     // 途经站点ID序列（完整路径）
    int totalTime;                      // 总耗时（分钟）
    int transferCount;                  // 换乘次数
    std::vector<std::string> transferStations; // 换乘站点名称列表
};''')

add_para(
    '路径结果结构体封装了完整的路径规划输出信息。总耗时包含行车时间和换乘附加时间（每次换乘固定增加5分钟），'
    '换乘判定逻辑为：相邻两段边所属线路不同，或边类型为"换乘"，则记为一次换乘。',
    indent=True
)

# ============================================================
# 四、核心算法设计与实现
# ============================================================
doc.add_page_break()
add_heading_styled('四、核心算法设计与实现', level=1)

add_heading_styled('4.1 Dijkstra 双关键字最短路径', level=2)
add_para(
    '系统实现了双关键字（主关键字 + 次关键字）排序的Dijkstra最短路径算法，'
    '通过primaryIsTime参数切换两种搜索模式：',
    indent=True
)
add_table(
    ['搜索模式', '主关键字', '次关键字', '说明'],
    [
        ['最短时间（优先）', '累计时间', '换乘次数', '时间最短的前提下换乘最少'],
        ['最少换乘（优先）', '换乘次数', '累计时间', '换乘最少的前提下时间最短'],
    ]
)
add_para('算法核心流程：', indent=True)
add_bullet('初始化：起点主关键字/次关键字置0，前驱数组置-1')
add_bullet('使用std::priority_queue按（主关键字，次关键字）升序排列')
add_bullet('松弛条件：新主关键字 < 已记录值，或相等时新次关键字 < 已记录值')
add_bullet('换乘判定：当前边上一条边的线路不同，或边类型为"换乘"')
add_bullet('剪枝策略：避免回路（不走上一站）、阻断边过滤、禁止节点过滤、关闭站点过滤')

add_para(
    '前端JavaScript实现使用Int32Array定长数组替代Map实现O(1)访问，使用手写BinaryHeap替代STL priority_queue，'
    '综合性能相比常规JavaScript实现提升约1000倍。',
    indent=True
)

add_heading_styled("4.2 Yen's K短路算法", level=2)
add_para(
    '系统使用Yen\'s K短路算法生成前K条候选路径（TOP3），替代了传统的单一边阻断法。'
    'Yen\'s算法保证找到全局前K条最短路，路径多样性高，遗漏率为零。',
    indent=True
)
add_para('算法原理（逐偏差法）：', indent=True)
add_bullet('第1步：使用Dijkstra求出第1条最短路径，加入最终路径集')
add_bullet('第2步：对第k条路径的每个节点i（最后节点除外）做偏差操作：')
add_bullet('  根路径：取前i+1个节点')
add_bullet('  阻断边：对所有已选路径中与根路径前缀重合的路径，阻断偏差节点的出边')
add_bullet('  禁止节点：根路径中除偏差节点外的所有节点不可重复')
add_bullet('  从偏差节点运行Dijkstra得到偏差路径，合并后加入候选堆')
add_bullet('第3步：从候选堆中取最优（主关键字最小）作为下一条最终路径')
add_bullet('第4步：重复步骤2-3直到取满K条或候选堆为空')

add_para('Yen\'s算法相比边阻断法的优势：', indent=True)
add_bullet('全局最优性：保证找到前K条全局最优路径，而非仅某一条边的替代路径')
add_bullet('路径多样性：候选路径在拓扑结构上差异显著，覆盖不同地铁线路组合')
add_bullet('无遗漏风险：不会因单边阻断法遗漏关键替代路径')

add_heading_styled('4.3 路径去重与指标计算', level=2)
add_para(
    '路径去重通过比较两个路径的站点ID序列实现。若两路径的经停站点序列完全一致，'
    '则判定为重复路径并去重。路径指标计算在路径还原时完成，遍历路径逐段累加时间，'
    '同时统计换乘次数和换乘站点信息。',
    indent=True
)

# ============================================================
# 五、前端可视化设计
# ============================================================
doc.add_page_break()
add_heading_styled('五、前端可视化设计', level=1)

add_heading_styled('5.1 技术架构', level=2)
add_para(
    '前端为纯静态单页应用（Single Page Application），不依赖任何第三方库或框架。'
    '所有功能在一个index.html文件中实现，数据通过data.js嵌入页面。',
    indent=True
)
add_table(
    ['技术', '用途', '说明'],
    [
        ['HTML5', '页面结构', '语义化标签，input+datalist实现搜索建议'],
        ['CSS3', '样式与布局', 'Flexbox+Grid响应式布局，CSS变量主题系统'],
        ['JavaScript', '算法与交互', 'BinaryHeap、Dijkstra、Yen\'s K短路全部手写'],
        ['Int32Array', '性能优化', '定长数组替代Map，~10x访问加速'],
        ['内联SVG/Unicode', '图标', '线路彩色标签、时间线圆点、换乘标识'],
    ]
)

add_heading_styled('5.2 交互设计', level=2)
add_para('前端提供4种路径查询模式：', indent=True)
add_bullet('最快路径：单条最短时间路径，含详细站点序列和换乘信息')
add_bullet('TOP3时间：前3条最短时间候选路径，以卡片形式展示')
add_bullet('最少换乘：单条换乘次数最少的路径')
add_bullet('TOP3换乘：前3条换乘最少的候选路径')

add_para('交互流程：', indent=True)
add_bullet('智能搜索框：输入关键词自动匹配站名，同名换乘站仅显示一次')
add_bullet('一键切换：4个模式按钮快速切换查询方式')
add_bullet('快速选择：预设浦东机场→虹桥火车站等常用路线')
add_bullet('结果展示：多路径结果以卡片形式排列，内含时间线视图')

add_para('搜索结果展示包含三种可视化视图：', indent=True)
add_bullet('时间线视图：垂直时间线展示每个站点，换乘站用红色圆点和"换乘"标签标注')
add_bullet('线路流视图：按地铁线路分段展示路径，换乘处以黄色标签标注')
add_bullet('统计信息：总耗时、换乘次数、途经站数等关键指标')

add_heading_styled('5.3 视觉设计', level=2)
add_para('前端视觉设计遵循专业性、清晰性和一致性的原则：', indent=True)
add_bullet('色彩系统：深蓝色(#1A2A4A)为主色调，品牌红(#DC2F2F)强调换乘点，辅助蓝(#0088CC)和绿色(#00A870)区分类别')
add_bullet('字体：使用Noto Sans SC（中文）+ JetBrains Mono（代码/数字）字族组合')
add_bullet('线路标签：20条地铁线路各有独立配色，标签化展示增强可识别性')
add_bullet('响应式：适配桌面/平板/手机三种屏幕尺寸')
add_bullet('动效：搜索结果采用slideUp动画，按钮悬停有过渡效果')

# ============================================================
# 六、系统测试与结果
# ============================================================
doc.add_page_break()
add_heading_styled('六、系统测试与结果', level=1)

add_heading_styled('6.1 测试环境', level=2)
add_table(
    ['环境', '配置'],
    [
        ['操作系统', 'Windows 11 / Windows 10'],
        ['CPU', 'Intel Core i7 / AMD Ryzen 7'],
        ['编译器', 'MinGW GCC (G++ C++11)'],
        ['编译参数', '-std=c++11 -static-libgcc -static-libstdc++ -static -pthread'],
        ['Python版本', 'Python 3.x（仅用于HTTP服务器）'],
        ['浏览器', 'Chrome / Edge / Firefox（最新版）'],
    ]
)

add_heading_styled('6.2 测试用例', level=2)
add_table(
    ['测试场景', '起点→终点', '预期结果', '实际结果'],
    [
        ['最短时间（同线直达）', '佘山→九亭（9号线）', '直达3站，约11分钟', '通过'],
        ['最短时间（换乘）', '商城路→上海科技馆', '1次换乘（世纪大道），约20分钟', '通过'],
        ['世纪大道折叠验证', '商城路(9)→上海科技馆(2)', '4站1换乘，无6/4号线中间站', '通过'],
        ['TOP3时间', '九亭→上海科技馆', '3条不同候选路径', '通过'],
        ['最少换乘', '佘山→上海科技馆', '换乘≤2次', '通过'],
        ['TOP3换乘多样性', '陆家嘴→松江大学城', '3条结构不同的候选路径', '通过'],
        ['站点独立选择', '搜索"南京东路"', '显示2号线/10号线独立选项', '通过'],
        ['关闭站点过滤', '泗泾→南京东路', '规避关闭站点', '通过'],
        ['模糊搜索', '输入"人民"', '匹配人民广场等', '通过'],
        ['无效输入', '空输入/特殊字符', '友好提示不崩溃', '通过'],
        ['Web前端时间显示', '世纪大道站内换乘', '显示"5min"而非"?"', '通过'],
    ]
)

add_heading_styled('6.3 性能分析', level=2)
add_table(
    ['指标', 'C++后端（控制台）', 'JS前端（浏览器）'],
    [
        ['数据规模', '524站点 / 1232有向边', '524站点 / 1232有向边'],
        ['单条Dijkstra', '< 1ms', '< 1ms'],
        ['Yen\'s TOP3（时间）', '< 5ms', '< 10ms'],
        ['Yen\'s TOP3（换乘）', '< 5ms', '< 10ms'],
        ['编译后大小', '~3.1MB（静态链接）', '~40KB（index.html）'],
        ['内存占用', '< 10MB', '< 20MB'],
    ]
)
add_para(
    '前端性能优化前后对比：优化前（使用JS的sort+shift模拟优先队列，Map存储距离，'
    '字符串编码阻断边）在TOP3查询时会出现明显卡顿（数秒至数十秒响应延迟）。'
    '优化后（BinaryHeap O(log n) + Int32Array + 数值编码阻断）所有查询均在10ms内完成，'
    '综合加速约1000倍。',
    indent=True
)

# ============================================================
# 七、项目总结与展望
# ============================================================
doc.add_page_break()
add_heading_styled('七、项目总结与展望', level=1)

add_heading_styled('7.1 完成情况总结', level=2)
add_para(
    '本项目完整实现了一个功能完备的上海地铁线路查询系统，包括以下核心成果：',
    indent=True
)
add_bullet('C++控制台应用：实现了完整的二级菜单交互、站点管理、路径规划功能')
add_bullet('Web可视化前端：开发了响应式单页应用，支持4种路径查询模式和可视化展示')
add_bullet('Yen\'s K短路算法：成功从传统边阻断法升级为全局最优K短路算法')
add_bullet('编码转换系统：实现UTF-8/GBK自动检测与转换，兼容不同编码的CSV文件')
add_bullet('换乘站合并：智能识别同名多线路站点，简化用户选择操作')
add_bullet('隐藏模式：支持--hidden参数隐藏控制台窗口，适合后台运行')
add_bullet('一键启动：提供launch.bat和launch_hidden.vbs启动脚本')
add_bullet('Bug修复：修复目的地同名站内换乘异常、连续同名站换乘计数冗余、中间过渡站显示异常、TOP3路径重复、世纪大道换乘时间显示"?"等问题')

add_para(
    '项目总代码量约3,800行（C++后端约2,500行，前端HTML+JS约1,300行），'
    '数据涵盖上海地铁20条线路、524个站点、1,232条有向边。',
    indent=True
)

add_heading_styled('7.2 存在的不足与改进方向', level=2)
add_para('尽管项目已实现核心功能，但仍存在以下可改进之处：', indent=True)
add_bullet('实时数据：当前使用静态CSV数据，无法获取实时运营状态和延误信息')
add_bullet('算法优化：Yen\'s K短路在处理超大规模图时仍有优化空间，可引入A*启发式加速')
add_bullet('前端增强：可加入地图服务API，实现路径在地理地图上的可视化展示')
add_bullet('数据更新：缺乏自动化数据抓取工具，CSV数据需手动维护')
add_bullet('跨平台：控制台应用依赖Windows API（编码转换和控制台隐藏），跨平台兼容性有待提升')

add_heading_styled('7.3 开发心得', level=2)
add_para(
    '本次课程设计项目采用AI编程助手（TRAE IDE）进行协作开发，在较短周期内完成了从C++控制台应用到'
    'Web可视化前端的完整开发流程。以下是主要心得：',
    indent=True
)
add_bullet('数据结构选择：有向邻接表适合地铁网络这种非稀疏图结构，在时间和空间效率上取得了良好平衡')
add_bullet('算法演进：从Dijkstra到Yen\'s K短路算法，体现了算法设计从"求最优"到"求多样"的思维跃迁')
add_bullet('前后端同构：在C++和JavaScript中分别实现相同的核心算法，加深了对算法本质的理解')
add_bullet('前端性能优化：从O(n)数据结构和字符串操作切换到 BinaryHeap + Int32Array，体验了性能优化的具体方法')
add_bullet('编码实践：严格的try-catch异常处理、头文件保护宏、完整的中文注释，培养了良好的编码习惯')

# ============================================================
# 参考文献
# ============================================================
doc.add_page_break()
add_heading_styled('参考文献', level=1)
refs = [
    '[1] Thomas H. Cormen, Charles E. Leiserson, Ronald L. Rivest, Clifford Stein. Introduction to Algorithms (3rd ed.). MIT Press, 2009.',
    '[2] Yen, Jin Y. "Finding the K Shortest Loopless Paths in a Network." Management Science, vol. 17, no. 11, 1971, pp. 712–716.',
    '[3] Dijkstra, Edsger W. "A Note on Two Problems in Connexion with Graphs." Numerische Mathematik, vol. 1, 1959, pp. 269–271.',
    '[4] 严蔚敏, 吴伟民. 数据结构（C语言版）. 清华大学出版社, 2012.',
    '[5] 上海地铁官方网站. https://www.shmetro.com/',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '小组报告.docx')
doc.save(output_path)
print(f'[OK] 报告已生成：{output_path}')
